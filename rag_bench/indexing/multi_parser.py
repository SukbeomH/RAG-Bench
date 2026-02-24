"""
멀티포맷 문서 파서 (multi_parser).

PDF / DOCX / HTML / TXT / MD 등 다양한 포맷의 문서를 Markdown 텍스트로 변환한다.
확장자를 자동 감지하여 적절한 파서를 선택한다.

외부 인터페이스:
  parse_document(path, *, sample, doc_type, max_chars) -> str
  parse_directory(dir_path, *, extensions, ...) -> list[tuple[Path, str]]

의존성:
  - PDF : pymupdf4llm (기존 설치됨)
  - DOCX: python-docx (pyproject.toml에 추가됨)
  - HTML: beautifulsoup4 + lxml (pyproject.toml에 추가됨)
  - TXT/MD: 내장 (의존성 없음)
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

# 지원하는 파일 확장자
SUPPORTED_EXTENSIONS = frozenset({
    ".pdf",
    ".docx", ".doc",
    ".html", ".htm",
    ".txt", ".md", ".text",
    ".csv",
})


# ---------------------------------------------------------------------------
# 포맷별 파서
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> str:
    """PDF → Markdown 변환 (pymupdf4llm)."""
    try:
        import pymupdf4llm
    except ImportError as e:
        raise ImportError("pymupdf4llm이 설치되어 있지 않습니다. pip install pymupdf4llm") from e

    md_text = pymupdf4llm.to_markdown(str(path))
    return md_text


def _parse_docx(path: Path) -> str:
    """DOCX → Markdown 변환 (python-docx)."""
    try:
        import docx
    except ImportError as e:
        raise ImportError("python-docx가 설치되어 있지 않습니다. pip install python-docx") from e

    doc = docx.Document(str(path))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 헤더 스타일 감지 → Markdown 헤딩 변환
        style_name = para.style.name if para.style else ""
        if "Heading 1" in style_name:
            parts.append(f"# {text}")
        elif "Heading 2" in style_name:
            parts.append(f"## {text}")
        elif "Heading 3" in style_name:
            parts.append(f"### {text}")
        else:
            parts.append(text)

    # 표 처리
    for table in doc.tables:
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
            if i == 0:
                rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
        parts.extend(rows)

    return "\n\n".join(parts)


def _parse_html(path: Path) -> str:
    """HTML → 순수 텍스트 변환 (beautifulsoup4)."""
    try:
        from bs4 import BeautifulSoup
    except ImportError as e:
        raise ImportError(
            "beautifulsoup4가 설치되어 있지 않습니다. pip install beautifulsoup4 lxml"
        ) from e

    html_content = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html_content, "lxml")

    # 스크립트/스타일 제거
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    # 헤딩 → Markdown 변환
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
        level = int(tag.name[1])
        tag.string = "#" * level + " " + (tag.get_text(strip=True) or "")

    # 순수 텍스트 추출 (공백 정규화)
    text = soup.get_text(separator="\n", strip=True)
    # 연속된 빈 줄 축소
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def _parse_html_string(html_content: str) -> str:
    """HTML 문자열 → 순수 텍스트 변환."""
    try:
        from bs4 import BeautifulSoup
    except ImportError as e:
        raise ImportError(
            "beautifulsoup4가 설치되어 있지 않습니다. pip install beautifulsoup4 lxml"
        ) from e

    soup = BeautifulSoup(html_content, "lxml")
    for tag in soup(["script", "style"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    return re.sub(r"\n{3,}", "\n\n", text)


def _parse_text(path: Path) -> str:
    """TXT / MD 파일 직접 읽기."""
    return path.read_text(encoding="utf-8", errors="ignore")


def _parse_csv(path: Path, max_rows: int = 500) -> str:
    """CSV → Markdown 표 변환."""
    try:
        import csv
        rows = []
        with open(path, newline="", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                if i >= max_rows:
                    rows.append(f"[... {i}행 이후 생략 ...]")
                    break
                rows.append("| " + " | ".join(str(cell) for cell in row) + " |")
                if i == 0:
                    rows.append("| " + " | ".join(["---"] * len(row)) + " |")
        return "\n".join(rows)
    except Exception as e:
        return f"CSV 파싱 오류: {e}"


# ---------------------------------------------------------------------------
# 공개 API
# ---------------------------------------------------------------------------

def parse_document(
    path: "str | Path",
    *,
    sample: bool = False,
    doc_type=None,
    max_chars: Optional[int] = None,
) -> str:
    """통합 문서 파서. 확장자를 자동 감지하여 적절한 파서를 선택한다.

    Args:
        path: 문서 파일 경로.
        sample: True이면 sampler.sample_text()로 샘플링 적용.
        doc_type: DocType 열거형 (sample=True일 때 사용).
        max_chars: 반환 최대 문자 수.

    Returns:
        파싱된 텍스트 (Markdown 또는 순수 텍스트).

    Raises:
        FileNotFoundError: 파일이 없을 때.
        ValueError: 지원하지 않는 파일 포맷일 때.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"파일 없음: {path}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        text = _parse_pdf(path)
    elif suffix in (".docx", ".doc"):
        text = _parse_docx(path)
    elif suffix in (".html", ".htm"):
        text = _parse_html(path)
    elif suffix in (".txt", ".md", ".text"):
        text = _parse_text(path)
    elif suffix == ".csv":
        text = _parse_csv(path)
    else:
        raise ValueError(
            f"지원하지 않는 파일 포맷: '{suffix}'. "
            f"지원 포맷: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    # 샘플링 적용
    if sample and doc_type is not None:
        from rag_bench.document_types.sampler import sample_text
        text = sample_text(text, doc_type, max_chars=max_chars or 50_000)
    elif max_chars is not None:
        text = text[:max_chars]

    return text


def parse_directory(
    dir_path: "str | Path",
    *,
    extensions: Optional[set[str]] = None,
    recursive: bool = False,
    sample: bool = False,
    doc_type=None,
    max_chars: Optional[int] = None,
    skip_errors: bool = True,
) -> list[tuple[Path, str]]:
    """디렉토리 내 문서를 일괄 파싱한다.

    Args:
        dir_path: 문서 디렉토리 경로.
        extensions: 처리할 확장자 집합 (None이면 SUPPORTED_EXTENSIONS 전체).
        recursive: True이면 하위 디렉토리까지 재귀 탐색.
        sample: True이면 샘플링 적용.
        doc_type: DocType (sample=True일 때 사용).
        max_chars: 파일당 최대 문자 수.
        skip_errors: True이면 파싱 오류 무시 (경고 출력).

    Returns:
        (Path, 텍스트) 튜플 목록.
    """
    dir_path = Path(dir_path)
    if not dir_path.is_dir():
        raise NotADirectoryError(f"디렉토리가 아님: {dir_path}")

    exts = extensions or SUPPORTED_EXTENSIONS
    pattern = "**/*" if recursive else "*"
    files = [f for f in dir_path.glob(pattern) if f.is_file() and f.suffix.lower() in exts]
    files.sort()

    results = []
    for f in files:
        try:
            text = parse_document(f, sample=sample, doc_type=doc_type, max_chars=max_chars)
            results.append((f, text))
            print(f"  파싱 완료: {f.name} ({len(text):,} chars)")
        except Exception as e:
            if skip_errors:
                print(f"  [경고] 파싱 실패: {f.name} — {e}")
            else:
                raise

    return results


def detect_extension_support() -> dict[str, bool]:
    """현재 환경에서 각 포맷 지원 여부를 반환한다 (의존성 체크)."""
    support = {}

    # PDF
    try:
        import pymupdf4llm  # noqa: F401
        support["pdf"] = True
    except ImportError:
        support["pdf"] = False

    # DOCX
    try:
        import docx  # noqa: F401
        support["docx"] = True
    except ImportError:
        support["docx"] = False

    # HTML
    try:
        from bs4 import BeautifulSoup  # noqa: F401
        support["html"] = True
    except ImportError:
        support["html"] = False

    # TXT/MD (항상 지원)
    support["txt"] = True
    support["md"] = True

    return support
